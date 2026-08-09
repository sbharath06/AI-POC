import os
from pathlib import Path
from PyPDF2 import PdfReader
from pdf2docx import Converter as PDFConverter
from docx import Document
from fpdf import FPDF
import openpyxl
import csv

CONVERTED_DIR = os.path.join(os.path.dirname(__file__), "..", "..", "converted")
os.makedirs(CONVERTED_DIR, exist_ok=True)

def convert_docx_to_pdf(input_path: str, output_path: str) -> str:
    """High-quality DOCX to PDF conversion using docx2pdf (MS Word COM) with ReportLab fallback."""
    extracted_text = ""
    
    # Method 1: docx2pdf (MS Word native automation - 100% pixel perfect)
    try:
        from docx2pdf import convert
        convert(input_path, output_path)
        if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
            doc = Document(input_path)
            extracted_text = "\n".join([p.text for p in doc.paragraphs if p.text])
            return extracted_text
    except Exception as e:
        print(f"[Converter] docx2pdf note/fallback: {e}")

    # Method 2: Rich ReportLab document builder (preserves headings, styles, tables, margins)
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        
        doc_docx = Document(input_path)
        extracted_text = "\n".join([p.text for p in doc_docx.paragraphs if p.text])

        pdf_doc = SimpleDocTemplate(
            output_path,
            pagesize=letter,
            rightMargin=54,
            leftMargin=54,
            topMargin=54,
            bottomMargin=54
        )

        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle(
            'DocxTitle',
            parent=styles['Heading1'],
            fontName='Helvetica-Bold',
            fontSize=22,
            leading=26,
            textColor=colors.HexColor('#1E293B'),
            spaceAfter=14
        )
        h1_style = ParagraphStyle(
            'DocxH1',
            parent=styles['Heading2'],
            fontName='Helvetica-Bold',
            fontSize=16,
            leading=20,
            textColor=colors.HexColor('#0F172A'),
            spaceBefore=12,
            spaceAfter=8
        )
        h2_style = ParagraphStyle(
            'DocxH2',
            parent=styles['Heading3'],
            fontName='Helvetica-Bold',
            fontSize=13,
            leading=16,
            textColor=colors.HexColor('#334155'),
            spaceBefore=10,
            spaceAfter=6
        )
        body_style = ParagraphStyle(
            'DocxBody',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=10.5,
            leading=14.5,
            textColor=colors.HexColor('#334155'),
            spaceAfter=6
        )

        story = []

        for p in doc_docx.paragraphs:
            text = p.text.strip()
            if not text:
                story.append(Spacer(1, 6))
                continue

            safe_text = (text.replace('&', '&amp;')
                             .replace('<', '&lt;')
                             .replace('>', '&gt;'))
            
            style_name = p.style.name.lower() if p.style else ""
            if 'title' in style_name:
                story.append(Paragraph(safe_text, title_style))
            elif 'heading 1' in style_name or style_name == 'h1':
                story.append(Paragraph(safe_text, h1_style))
            elif 'heading 2' in style_name or style_name == 'h2' or 'heading 3' in style_name:
                story.append(Paragraph(safe_text, h2_style))
            else:
                story.append(Paragraph(safe_text, body_style))

        # Process docx tables
        for table in doc_docx.tables:
            table_data = []
            for row in table.rows:
                row_data = [Paragraph(cell.text.strip().replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;'), body_style) for cell in row.cells]
                table_data.append(row_data)
            if table_data:
                t = Table(table_data)
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#F1F5F9')),
                    ('TEXTCOLOR', (0,0), (-1,0), colors.HexColor('#0F172A')),
                    ('ALIGN', (0,0), (-1,-1), 'LEFT'),
                    ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                    ('BOTTOMPADDING', (0,0), (-1,-1), 6),
                    ('TOPPADDING', (0,0), (-1,-1), 6),
                    ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#CBD5E1')),
                ]))
                story.append(Spacer(1, 8))
                story.append(t)
                story.append(Spacer(1, 8))

        pdf_doc.build(story)
        return extracted_text

    except Exception as e:
        print(f"[Converter] ReportLab fallback error: {e}")

    # Method 3: Fallback plain text FPDF
    doc = Document(input_path)
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Helvetica", size=11)
    page_width = pdf.w - pdf.l_margin - pdf.r_margin
    for para in doc.paragraphs:
        extracted_text += para.text + "\n"
        text = para.text.strip()
        if not text:
            pdf.ln(4)
            continue
        text = text.encode('latin-1', 'replace').decode('latin-1')
        pdf.multi_cell(page_width, 7, text)
    pdf.output(output_path)
    return extracted_text


def convert_file(input_path: str, output_format: str) -> tuple[str, str]:
    input_path = os.path.abspath(input_path)
    base_name = os.path.splitext(os.path.basename(input_path))[0]
    output_filename = f"{base_name}.{output_format}"
    output_path = os.path.join(CONVERTED_DIR, output_filename)
    
    extracted_text = ""
    ext = input_path.split('.')[-1].lower()
    
    try:
        if ext == 'pdf' and output_format == 'txt':
            reader = PdfReader(input_path)
            for page in reader.pages:
                extracted_text += page.extract_text() + "\n"
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(extracted_text)
                
        elif ext == 'pdf' and output_format == 'docx':
            cv = PDFConverter(input_path)
            cv.convert(output_path)
            cv.close()
            reader = PdfReader(input_path)
            for page in reader.pages:
                extracted_text += page.extract_text() + "\n"
                
        elif ext == 'docx' and output_format == 'pdf':
            extracted_text = convert_docx_to_pdf(input_path, output_path)

            
        elif ext == 'docx' and output_format == 'txt':
            doc = Document(input_path)
            for para in doc.paragraphs:
                extracted_text += para.text + "\n"
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(extracted_text)
                
        elif ext == 'txt' and output_format == 'pdf':
            with open(input_path, 'r', encoding='utf-8') as f:
                extracted_text = f.read()
            pdf = FPDF()
            pdf.add_page()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.set_font("Helvetica", size=11)
            text = extracted_text.encode('latin-1', 'replace').decode('latin-1')
            pdf.multi_cell(0, 7, text)
            pdf.output(output_path)
            
        elif ext == 'md' and output_format == 'pdf':
            with open(input_path, 'r', encoding='utf-8') as f:
                extracted_text = f.read()
            # Strip simple markdown
            text_to_print = extracted_text.replace('#', '').replace('*', '')
            pdf = FPDF()
            pdf.add_page()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.set_font("Helvetica", size=11)
            text = text_to_print.encode('latin-1', 'replace').decode('latin-1')
            pdf.multi_cell(0, 7, text)
            pdf.output(output_path)
            
        elif ext == 'xlsx' and output_format == 'pdf':
            wb = openpyxl.load_workbook(input_path, data_only=True)
            sheet = wb.active
            pdf = FPDF()
            pdf.add_page()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.set_font("Helvetica", size=11)
            for row in sheet.iter_rows(values_only=True):
                row_str = " | ".join([str(cell) if cell is not None else "" for cell in row])
                extracted_text += row_str + "\n"
                text = row_str.encode('latin-1', 'replace').decode('latin-1')
                pdf.multi_cell(0, 7, text)
            pdf.output(output_path)
            
        elif ext == 'xlsx' and output_format == 'csv':
            wb = openpyxl.load_workbook(input_path, data_only=True)
            sheet = wb.active
            with open(output_path, 'w', encoding='utf-8', newline='') as f:
                writer = csv.writer(f)
                for row in sheet.iter_rows(values_only=True):
                    writer.writerow(row)
                    row_str = " | ".join([str(cell) if cell is not None else "" for cell in row])
                    extracted_text += row_str + "\n"
                    
        elif ext == 'csv' and output_format == 'pdf':
            pdf = FPDF()
            pdf.add_page()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.set_font("Helvetica", size=11)
            with open(input_path, 'r', encoding='utf-8') as f:
                reader = csv.reader(f)
                for row in reader:
                    row_str = " | ".join(row)
                    extracted_text += row_str + "\n"
                    text = row_str.encode('latin-1', 'replace').decode('latin-1')
                    pdf.multi_cell(0, 7, text)
            pdf.output(output_path)
        else:
            raise ValueError(f"Unsupported conversion: {ext} -> {output_format}")
            
    except Exception as e:
        raise Exception(f"Conversion failed: {str(e)}")
        
    return output_path, extracted_text
